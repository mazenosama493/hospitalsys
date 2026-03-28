"""
Generate the VirtualHospital backend specification DOCX.
Run: /opt/homebrew/bin/python3 generate_backend_spec.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = RGBColor(0x0D, 0x47, 0x8A) if level == 1 else (
        RGBColor(0x15, 0x65, 0xC0) if level == 2 else RGBColor(0x19, 0x76, 0xD2))
    return h

def add_para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, "1565C0")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "EEF2FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "90CAF9")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_note(doc, text, color="FFF9C4"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    run = p.add_run("ℹ  " + text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x00)
    return p

# ─────────────────────────────────────────────────────────────────────────────
# Document Setup
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default paragraph font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(60)
run = title_para.add_run("VirtualHospital (MedHub)")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x0D, 0x47, 0x8A)

sub1 = doc.add_paragraph()
sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub1.add_run("Backend API & Data Model Specification")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run(f"Version 1.0  ·  {datetime.date.today().strftime('%B %d, %Y')}")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
sub3 = doc.add_paragraph()
sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub3.add_run("Prepared for: Backend Development Team")
run.font.size = Pt(11)
run.italic = True

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS (manual)
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "Table of Contents", level=1)
toc_items = [
    ("1", "Project Overview"),
    ("2", "Technology Stack & Conventions"),
    ("3", "Authentication & Authorization"),
    ("4", "Global Shared Types"),
    ("5", "Module 1 — Front Desk / ADT Portal"),
    ("6", "Module 2 — Doctor Portal"),
    ("7", "Module 3 — Nurse Portal"),
    ("8", "Module 4 — Laboratory (LIS) Portal"),
    ("9", "Module 5 — Radiology (RIS/PACS) Portal"),
    ("10", "Module 6 — Pharmacy Portal"),
    ("11", "Module 7 — Billing / Revenue Cycle Portal"),
    ("12", "Module 8 — Admin Portal"),
    ("13", "Module 9 — CDSS (Clinical Decision Support)"),
    ("14", "Module 10 — Patient Portal"),
    ("15", "Real-Time & Notification Requirements"),
    ("16", "Pagination, Filtering & Search Conventions"),
    ("17", "Error Response Format"),
    ("18", "Audit Logging Requirements"),
    ("19", "File / Document Upload"),
    ("20", "Summary Endpoint Checklist"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f"  {num}.  {title}")
    run.font.size = Pt(10)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 — PROJECT OVERVIEW
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "1. Project Overview")
add_para(doc,
    "MedHub is a comprehensive, multi-role virtual hospital management system. "
    "The frontend is a Next.js 16 application with a vertical-slice architecture. "
    "Every clinical department has its own portal, all sharing a common authentication layer and REST/WebSocket backend.",
    size=10)
doc.add_paragraph()

add_heading(doc, "1.1 User Roles", level=2)
add_table(doc,
    ["Role Key", "Display Name", "Default Route", "Description"],
    [
        ["admin",         "Administrator", "/admin",      "System configuration, user management, audit"],
        ["doctor",        "Doctor",        "/doctor",     "Clinical care, SOAP notes, orders, prescriptions"],
        ["nurse",         "Nurse",         "/nurse",      "Bedside care, vitals, MAR, nursing notes, handoffs"],
        ["lab_tech",      "Lab Technician","/lab",        "Specimen processing, result entry, critical notifications"],
        ["radiologist",   "Radiologist",   "/radiology",  "Imaging orders, PACS, reporting, critical findings"],
        ["pharmacist",    "Pharmacist",    "/pharmacy",   "Rx verification, dispensing, drug safety, formulary"],
        ["billing_staff", "Billing Staff", "/billing",    "Invoices, insurance claims, payments, denials"],
        ["front_desk",    "Front Desk",    "/frontdesk",  "Patient registration, ADT, queue management"],
        ["patient",       "Patient",       "/patient",    "Self-service portal: appointments, records, Rx"],
    ],
    col_widths=[1.1, 1.3, 1.5, 3.2])

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 — TECH STACK & CONVENTIONS
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "2. Technology Stack & Conventions")

add_heading(doc, "2.1 Frontend Stack", level=2)
add_table(doc,
    ["Layer", "Technology"],
    [
        ["Framework",    "Next.js 16 (App Router) + React 19"],
        ["Language",     "TypeScript 5 (strict mode)"],
        ["State",        "Zustand 5 (per-module stores, persisted auth)"],
        ["Data Fetching","TanStack Query v5 (@tanstack/react-query)"],
        ["Forms",        "react-hook-form v7 + Zod v4 validation"],
        ["UI",           "shadcn/ui (Radix UI primitives) + Tailwind CSS 4"],
        ["Icons",        "Lucide React 0.577"],
        ["Date",         "date-fns v4"],
    ],
    col_widths=[1.8, 5.4])

doc.add_paragraph()
add_heading(doc, "2.2 API Conventions", level=2)
add_bullet(doc, "Base URL: https://api.medhub.hospital/v1")
add_bullet(doc, "All request/response bodies: application/json")
add_bullet(doc, "Timestamps: ISO 8601 strings (e.g. 2026-03-16T14:35:00Z)")
add_bullet(doc, "IDs: string UUIDs (uuid v4)")
add_bullet(doc, "Pagination: cursor-based or page/limit (see Section 16)")
add_bullet(doc, "Authentication: Bearer JWT in Authorization header")
add_bullet(doc, "Role filtering: backend MUST enforce role-based access — never trust the frontend to filter sensitive data")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3 — AUTHENTICATION & AUTHORIZATION
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "3. Authentication & Authorization")

add_heading(doc, "3.1 Auth Flow", level=2)
add_para(doc,
    "The frontend uses a role-selection login form. The user picks their role, enters email + password. "
    "The backend validates credentials, confirms the requested role matches the stored role, and returns a JWT.")

add_heading(doc, "3.2 Endpoints", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["POST", "/auth/login",         "Authenticate. Body: { email, password, role }. Returns: { token, user }"],
        ["POST", "/auth/logout",        "Invalidate token / session"],
        ["POST", "/auth/refresh",       "Refresh JWT. Body: { refreshToken }"],
        ["GET",  "/auth/me",            "Return current authenticated user profile"],
        ["PUT",  "/auth/me/password",   "Change own password. Body: { currentPassword, newPassword }"],
    ],
    col_widths=[0.7, 2.2, 4.3])

add_heading(doc, "3.3 Login Request / Response Shape", level=2)
add_code(doc, "POST /auth/login")
add_code(doc, "Request:  { email: string, password: string, role: UserRole }")
add_code(doc, "Response: {")
add_code(doc, "  token:        string,   // JWT access token (exp: 8h)")
add_code(doc, "  refreshToken: string,   // exp: 30 days")
add_code(doc, "  user: {")
add_code(doc, "    id: string, email: string, firstName: string, lastName: string,")
add_code(doc, "    role: UserRole, department?: string, avatar?: string")
add_code(doc, "  }")
add_code(doc, "}")

add_heading(doc, "3.4 JWT Payload", level=2)
add_code(doc, "{ sub: userId, role: UserRole, departmentId?: string, iat, exp }")

add_heading(doc, "3.5 Role-Based Access Control", level=2)
add_para(doc, "Each endpoint must be protected so only the appropriate roles can access it. "
    "Suggested middleware: check JWT role against an allowed-roles list per route. "
    "Unauthorised access must return HTTP 403 with error code FORBIDDEN.")
add_note(doc, "The frontend DOES NOT filter data by role — it trusts the API to return only what that role is entitled to see.")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 — GLOBAL SHARED TYPES
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "4. Global Shared Types")
add_para(doc, "The following enumerations and models are shared across multiple modules.")

add_heading(doc, "4.1 Priority", level=2)
add_code(doc, "type Priority = 'low' | 'normal' | 'high' | 'urgent' | 'stat'")
add_para(doc, "Used by orders, queue entries, imaging studies, specimens, etc.")

add_heading(doc, "4.2 PatientStatus", level=2)
add_code(doc, "type PatientStatus = 'active' | 'discharged' | 'critical' | 'stable' | 'admitted'")

add_heading(doc, "4.3 Patient (core object)", level=2)
add_table(doc,
    ["Field", "Type", "Required", "Notes"],
    [
        ["id",               "string (UUID)", "✓", "System-generated"],
        ["mrn",              "string",        "✓", "Medical Record Number — system-generated, unique"],
        ["firstName",        "string",        "✓", ""],
        ["lastName",         "string",        "✓", ""],
        ["dateOfBirth",      "string",        "✓", "ISO 8601 date: YYYY-MM-DD"],
        ["gender",           "male|female|other","✓",""],
        ["phone",            "string",        "✓", ""],
        ["email",            "string",        "✓", ""],
        ["address",          "string",        "✓", ""],
        ["bloodType",        "string",        "",  "e.g. A+, O-"],
        ["allergies",        "string[]",      "",  "Free-text allergy list"],
        ["status",           "PatientStatus", "✓", ""],
        ["insuranceProvider","string",        "",  ""],
        ["insuranceId",      "string",        "",  ""],
        ["admissionDate",    "string",        "",  "ISO 8601"],
        ["assignedDoctor",   "string",        "",  "Doctor name or ID"],
        ["ward",             "string",        "",  ""],
        ["roomNumber",       "string",        "",  ""],
    ],
    col_widths=[1.4, 1.4, 0.7, 3.7])

add_heading(doc, "4.4 Appointment", level=2)
add_table(doc,
    ["Field", "Type", "Required", "Notes"],
    [
        ["id",         "string", "✓", ""],
        ["patientId",  "string", "✓", "FK → Patient"],
        ["patientName","string", "✓", "Denormalised for display"],
        ["doctorId",   "string", "✓", "FK → User (doctor)"],
        ["doctorName", "string", "✓", ""],
        ["department", "string", "✓", ""],
        ["date",       "string", "✓", "YYYY-MM-DD"],
        ["time",       "string", "✓", "HH:mm"],
        ["duration",   "number", "✓", "Minutes"],
        ["status",     "AppointmentStatus","✓","scheduled|in-progress|completed|cancelled|no-show"],
        ["type",       "string", "✓", "consultation|follow-up|procedure|telemedicine"],
        ["notes",      "string", "",  ""],
    ],
    col_widths=[1.3, 1.2, 0.7, 4.0])

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5 — FRONT DESK / ADT
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "5. Module 1 — Front Desk / ADT Portal")
add_para(doc, "Accessible by: front_desk (primary), admins. Covers patient registration, "
    "admission/discharge/transfer (ADT), bed management, queue management, and consent documents.")

add_heading(doc, "5.1 Patient Registration Endpoints", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",    "/patients",                  "List patients. Filter: status, search (name/MRN/phone)"],
        ["GET",    "/patients/:id",              "Get single patient (full ADTPatient object)"],
        ["POST",   "/patients",                  "Register new patient. Returns created patient with system MRN"],
        ["PUT",    "/patients/:id",              "Update patient demographics"],
        ["GET",    "/patients/search",           "Quick search by name, MRN, phone, DOB. Query: q="],
        ["GET",    "/patients/:id/duplicates",   "Check for potential duplicate records (DuplicateCandidate[])"],
        ["POST",   "/patients/merge",            "Merge two duplicate records. Body: { keepId, mergeId }"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "5.2 ADTPatient — Full Registration Shape", level=2)
add_para(doc, "Extends Patient with additional fields required at registration:")
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["emergencyContact",    "EmergencyContact",  "{ name, relationship, phone }"],
        ["nationality",         "string",            ""],
        ["maritalStatus",       "string",            "single|married|divorced|widowed"],
        ["preferredLanguage",   "string",            ""],
        ["consentSigned",       "boolean",           "General consent signed at admission"],
        ["registeredAt",        "string",            "ISO 8601 datetime"],
        ["insurance",           "Insurance",         "{ id, provider, policyNumber, groupNumber?, validFrom, validTo, copay?, coverageType }"],
    ],
    col_widths=[1.8, 1.6, 3.8])

add_heading(doc, "5.3 Admission / Discharge / Transfer", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/admissions",           "List admissions. Filter: status, department, ward, date range"],
        ["GET",  "/admissions/:id",       "Single admission"],
        ["POST", "/admissions",           "Create admission. Body: AdmissionFormData"],
        ["PUT",  "/admissions/:id/status","Update status. Body: { status: AdmissionStatus }"],
        ["POST", "/admissions/:id/discharge","Discharge patient. Body: DischargeFormData"],
        ["POST", "/admissions/:id/transfer", "Transfer patient. Body: TransferFormData"],
        ["GET",  "/beds",                 "List all beds. Filter: ward, status, type"],
        ["GET",  "/beds/:bedId",          "Single bed"],
        ["PUT",  "/beds/:bedId/status",   "Update bed status. Body: { status: BedStatus }"],
        ["GET",  "/wards",               "List all wards"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "5.4 Queue Management", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/queue",               "List queue entries. Filter: service, status, date"],
        ["POST", "/queue",               "Add patient to queue. Body: { patientId, service, priority }"],
        ["PUT",  "/queue/:id/status",    "Update status (called, serving, completed, no-show)"],
        ["GET",  "/queue/stats",         "Current wait counts per service window"],
    ],
    col_widths=[0.7, 2.0, 4.5])

add_heading(doc, "5.5 Consent Documents", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/:id/consents",   "List consent documents for patient"],
        ["POST", "/patients/:id/consents",   "Create consent document record"],
        ["PUT",  "/consents/:id/sign",       "Mark consent as signed. Body: { signedBy, fileUrl? }"],
    ],
    col_widths=[0.7, 2.5, 4.0])

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6 — DOCTOR PORTAL
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "6. Module 2 — Doctor Portal")
add_para(doc, "Accessible by: doctor. Covers schedule, patient chart (SOAP encounters), "
    "orders, prescriptions, referrals, result review, and CDSS integration.")

add_heading(doc, "6.1 Schedule", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/appointments",                   "List appointments. Filter: doctorId, date, status"],
        ["GET",  "/appointments/:id",               "Single appointment"],
        ["POST", "/appointments",                   "Book appointment"],
        ["PUT",  "/appointments/:id",               "Update appointment"],
        ["PUT",  "/appointments/:id/status",        "Update status (arrived, cancelled, completed, etc.)"],
        ["DELETE","/appointments/:id",              "Cancel appointment"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "6.2 Patient Chart — Encounters (SOAP Notes)", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/:id/encounters",    "List all encounters for patient"],
        ["GET",  "/encounters/:id",             "Single encounter (full SOAP)"],
        ["POST", "/encounters",                 "Create new encounter. Body: EncounterData"],
        ["PUT",  "/encounters/:id",             "Update encounter (autosave)"],
        ["POST", "/encounters/:id/sign",        "Cosign/lock encounter. Sets status=signed, signedAt"],
        ["POST", "/encounters/:id/amend",       "Amend signed encounter (creates an amendment record)"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "6.3 Encounter Object", level=2)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",          "string",          ""],
        ["patientId",   "string",          "FK → Patient"],
        ["patientName", "string",          "Denormalised"],
        ["date",        "string",          "YYYY-MM-DD"],
        ["subjective",  "string",          "Patient's reported symptoms"],
        ["objective",   "string",          "Exam findings, vitals, labs"],
        ["assessment",  "string",          "Diagnoses / clinical impression"],
        ["plan",        "string",          "Treatment plan, follow-up"],
        ["status",      "EncounterStatus", "in-progress|completed|signed|amended"],
        ["authorId",    "string",          "FK → User (doctor)"],
        ["authorName",  "string",          ""],
        ["signedAt",    "string?",         "ISO 8601, present when status=signed"],
        ["visitType",   "string",          "inpatient|outpatient|emergency|telemedicine"],
    ],
    col_widths=[1.3, 1.5, 4.4])

add_heading(doc, "6.4 Orders", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/:id/orders",    "List orders for patient. Filter: category, status"],
        ["GET",  "/orders/:id",             "Single order with results if available"],
        ["POST", "/orders",                 "Place order. Body: OrderItem (without id/orderedAt/status)"],
        ["PUT",  "/orders/:id/status",      "Update order status. Body: { status, completedAt?, results? }"],
        ["DELETE","/orders/:id",            "Cancel pending order"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "6.5 Prescriptions", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/:id/prescriptions", "List prescriptions for patient. Filter: status"],
        ["GET",  "/prescriptions/:id",          "Single prescription"],
        ["POST", "/prescriptions",              "Write prescription. Body: PrescriptionData"],
        ["PUT",  "/prescriptions/:id",          "Update prescription"],
        ["PUT",  "/prescriptions/:id/status",   "Discontinue / put on hold. Body: { status }"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "6.6 Diagnoses", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/:id/diagnoses",   "List diagnoses for patient"],
        ["POST", "/diagnoses",                "Add diagnosis. Body: includes ICD-10 code"],
        ["PUT",  "/diagnoses/:id",            "Update diagnosis status"],
        ["GET",  "/icd10/search",             "ICD-10 typeahead. Query: q="],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "6.7 Referrals", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/:id/referrals",  "List referrals for patient"],
        ["POST", "/referrals",               "Create referral"],
        ["PUT",  "/referrals/:id/status",    "Accept / complete / decline referral"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "6.8 Results Inbox", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/doctors/:id/results",          "Pending results for doctor review. Filter: category (lab|imaging), flag"],
        ["PUT",  "/results/:id/review",           "Mark result as reviewed. Body: { reviewedBy, notes? }"],
    ],
    col_widths=[0.7, 2.6, 3.9])

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 7 — NURSE PORTAL
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "7. Module 3 — Nurse Portal")
add_para(doc, "Accessible by: nurse. Covers ward census, vitals flowsheet, intake/output, "
    "medication administration record (MAR), nursing notes, tasks, pain assessment, "
    "wound care, handoffs, and discharge checklists.")

add_heading(doc, "7.1 Ward Census", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET", "/wards/:wardId/patients",   "List all patients in a ward with bed info, vitals summary, task counts"],
        ["GET", "/wards",                    "List accessible wards"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "7.2 Vitals Flowsheet", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/:id/vitals",     "List vitals entries. Filter: dateFrom, dateTo, limit"],
        ["POST", "/patients/:id/vitals",     "Record new vitals entry. Body: VitalEntry (without id)"],
        ["GET",  "/patients/:id/vitals/latest", "Most recent vitals set"],
    ],
    col_widths=[0.7, 2.5, 4.0])
add_heading(doc, "VitalEntry Object", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",               "string",  ""],
        ["patientId",        "string",  ""],
        ["patientName",      "string",  ""],
        ["timestamp",        "string",  "ISO 8601"],
        ["systolic",         "number",  "mmHg"],
        ["diastolic",        "number",  "mmHg"],
        ["heartRate",        "number",  "bpm"],
        ["temperature",      "number",  "°F"],
        ["spo2",             "number",  "% oxygen saturation"],
        ["respiratoryRate",  "number",  "breaths/min"],
        ["painScore",        "number?", "0–10"],
        ["gcs",              "number?", "Glasgow Coma Scale 3–15"],
        ["recordedBy",       "string",  "Nurse name/ID"],
        ["notes",            "string?", ""],
    ],
    col_widths=[1.5, 1.2, 4.5])

add_heading(doc, "7.3 Intake & Output", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/:id/io",      "IO entries. Filter: date, direction"],
        ["POST", "/patients/:id/io",      "Record IO entry. Body: IntakeOutput (without id)"],
        ["DELETE","/io/:id",              "Delete erroneous entry"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "7.4 Pain Assessments", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/:id/pain",    "Pain assessment history"],
        ["POST", "/patients/:id/pain",    "Record pain assessment. Body: PainEntry (without id)"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "7.5 Medication Administration Record (MAR)", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/:id/mar",               "MAR entries. Filter: date, status"],
        ["PUT",  "/mar/:id/administer",             "Document drug administration. Body: { administeredTime, administeredBy, barcode?, notes? }"],
        ["PUT",  "/mar/:id/status",                 "Mark held, missed, refused. Body: { status, notes? }"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "7.6 Nursing Notes", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/:id/nursing-notes",   "List nursing notes"],
        ["POST", "/patients/:id/nursing-notes",   "Create note. Body: { category, content }"],
        ["PUT",  "/nursing-notes/:id",            "Edit note (within allowed window)"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "7.7 Nursing Tasks", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/nurses/:id/tasks",             "Task list. Filter: status, wardId, shift"],
        ["GET",  "/patients/:id/tasks",           "Tasks for specific patient"],
        ["POST", "/tasks",                        "Create task"],
        ["PUT",  "/tasks/:id/complete",           "Mark complete. Body: { completedBy, completedTime, notes? }"],
        ["PUT",  "/tasks/:id/status",             "Update status"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "7.8 Wound Notes", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/:id/wounds",    "Wound documentation"],
        ["POST", "/patients/:id/wounds",    "Add wound note"],
        ["PUT",  "/wounds/:id",             "Update wound note"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "7.9 Handoffs", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/handoffs",             "Handoffs. Filter: wardId, shiftDate, shiftType"],
        ["POST", "/handoffs",             "Create SBAR handoff. Body: HandoffEntry (without id)"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "7.10 Discharge Checklist", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/:id/discharge-checklist",   "Get checklist items"],
        ["PUT",  "/discharge-checklist/:id/complete",   "Mark item complete. Body: { completedBy, notes? }"],
        ["POST", "/patients/:id/discharge-checklist",   "Generate default checklist from template"],
    ],
    col_widths=[0.7, 2.7, 3.8])

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 8 — LABORATORY (LIS)
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "8. Module 4 — Laboratory (LIS) Portal")
add_para(doc, "Accessible by: lab_tech (primary), doctor (results view). Covers specimen management, "
    "accession, analyzer queue, result entry, verification, critical value notification, "
    "and lab reports.")

add_heading(doc, "8.1 Worklist", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET", "/lab/worklist",   "Lab order worklist. Filter: status (pending, in-progress, completed, stat), priority, date"],
    ],
    col_widths=[0.7, 2.0, 4.5])

add_heading(doc, "8.2 Specimens", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/specimens",               "List specimens. Filter: status, type, date"],
        ["GET",  "/specimens/:id",           "Single specimen"],
        ["PUT",  "/specimens/:id/status",    "Update specimen status/condition"],
        ["POST", "/specimens/:id/reject",    "Reject specimen. Body: { reason, notes? }"],
        ["POST", "/specimens/recollect",     "Create recollection request. Body: RecollectionRequest"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "Specimen Object", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",            "string",          ""],
        ["barcode",       "string",          "Unique barcode / label ID"],
        ["patientId",     "string",          "FK → Patient"],
        ["patientName",   "string",          "Denormalised"],
        ["type",          "SpecimenType",    "blood|serum|plasma|urine|csf|stool|swab|tissue|other"],
        ["collectionTime","string?",         "ISO 8601"],
        ["collectedBy",   "string?",         ""],
        ["receivedAt",    "string?",         "ISO 8601"],
        ["status",        "SpecimenStatus",  "ordered→collected→in-transit→received→processing→analyzed→resulted|rejected"],
        ["condition",     "SpecimenCondition","acceptable|hemolyzed|lipemic|icteric|clotted|insufficient|wrong-tube"],
        ["orderId",       "string",          "FK → OrderItem"],
        ["testNames",     "string[]",         ""],
        ["notes",         "string?",         ""],
    ],
    col_widths=[1.4, 1.5, 4.3])

add_heading(doc, "8.3 Accession", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/accessions",        "List accession records"],
        ["POST", "/accessions",        "Accession a specimen. Auto-generates accession number"],
        ["GET",  "/accessions/:id",    "Single accession record"],
    ],
    col_widths=[0.7, 2.0, 4.5])

add_heading(doc, "8.4 Analyzer Queue", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/analyzers/queue",              "Current analyzer queue. Filter: instrument, status"],
        ["PUT",  "/analyzers/queue/:id/status",   "Update run status"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "8.5 Result Entry & Verification", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/lab/panels",              "List panels. Filter: status, patientId"],
        ["GET",  "/lab/panels/:id",          "Panel with all test results"],
        ["POST", "/lab/panels/:id/results",  "Enter or update results for a panel"],
        ["PUT",  "/lab/panels/:id/verify",   "Verify / release results. Body: { verifiedBy }"],
        ["GET",  "/lab/results/:id",         "Single LabTestResult"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "LabTestResult Object", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",             "string",       ""],
        ["specimenId",     "string",       "FK → Specimen"],
        ["testCode",       "string",       "LOINC code"],
        ["testName",       "string",       ""],
        ["value",          "string",       "Stored as string (allows text results)"],
        ["unit",           "string",       "e.g. g/dL, mmol/L"],
        ["referenceRange", "string",       "e.g. 3.5–5.0"],
        ["flag",           "LabResultFlag","normal|high|low|critical-high|critical-low"],
        ["previousValue",  "string?",      "For delta check display"],
        ["delta",          "string?",      "e.g. +2.1 g/dL"],
        ["method",         "string?",      "Analyzer method used"],
        ["analyzedAt",     "string?",      "ISO 8601"],
        ["verifiedBy",     "string?",      ""],
        ["verifiedAt",     "string?",      "ISO 8601"],
        ["status",         "string",       "pending|preliminary|final|corrected|cancelled"],
    ],
    col_widths=[1.4, 1.3, 4.5])

add_heading(doc, "8.6 Lab Reports", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/lab/reports",         "List lab reports. Filter: patientId, status, date"],
        ["GET",  "/lab/reports/:id",     "Single report with all results"],
        ["PUT",  "/lab/reports/:id/release", "Release report to ordering clinician"],
    ],
    col_widths=[0.7, 2.2, 4.3])

add_heading(doc, "8.7 Critical Values", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/lab/critical",                          "Unacknowledged critical results"],
        ["POST", "/lab/critical/:resultId/notify",         "Record notification attempt. Body: { notifiedTo, callbackTime? }"],
        ["PUT",  "/lab/critical/:resultId/acknowledge",    "Clinician-acknowledged. Body: { acknowledgedBy }"],
    ],
    col_widths=[0.7, 2.7, 3.8])
add_note(doc, "Critical value notifications must also be sent via WebSocket event: lab.critical_result")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 9 — RADIOLOGY (RIS/PACS)
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "9. Module 5 — Radiology (RIS/PACS) Portal")
add_para(doc, "Accessible by: radiologist (primary), doctor (results). Covers imaging orders, "
    "protocoling, scheduling, study acquisition, PACS viewing, reporting, critical findings, "
    "and prior study comparison.")

add_heading(doc, "9.1 Imaging Orders", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/radiology/orders",              "Imaging order worklist. Filter: status, modality, priority, date"],
        ["GET",  "/radiology/orders/:id",          "Single order"],
        ["POST", "/radiology/orders",              "Place imaging order (from doctor portal)"],
        ["PUT",  "/radiology/orders/:id/protocol", "Protocol order. Body: { protocoledBy, notes? }"],
        ["PUT",  "/radiology/orders/:id/schedule", "Schedule study. Body: { scheduledAt, room, technologist }"],
        ["PUT",  "/radiology/orders/:id/assign",   "Assign to radiologist. Body: { radiologistId }"],
        ["PUT",  "/radiology/orders/:id/cancel",   "Cancel order"],
    ],
    col_widths=[0.7, 2.7, 3.8])

add_heading(doc, "ImagingOrder Object — Key Fields", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",              "string",            ""],
        ["accessionNumber", "string",            "System-generated (RIS format: MODALITY-YYYYMMDD-####)"],
        ["patientId",       "string",            ""],
        ["modality",        "ImagingModality",   "XR|CT|MRI|US|NM|PET|DEXA|FLUORO|MAMMO"],
        ["examCode",        "string",            "CPT / local exam code"],
        ["examName",        "string",            "Human-readable exam name"],
        ["bodyRegion",      "string",            ""],
        ["laterality",      "string?",           "left|right|bilateral"],
        ["contrastRequired","boolean",           ""],
        ["priority",        "Priority",          ""],
        ["clinicalHistory", "string",            "Required input from ordering physician"],
        ["status",          "ImagingStudyStatus","ordered→protocoled→scheduled→arrived→in-progress→acquired→reading→reported→signed"],
    ],
    col_widths=[1.5, 1.5, 4.2])

add_heading(doc, "9.2 Studies & PACS", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/radiology/studies",           "List studies. Filter: modality, status, date, radiologistId"],
        ["GET",  "/radiology/studies/:id",       "Study detail including PACS URL + report link"],
        ["POST", "/radiology/studies",           "Create study record (linked to order, set by technologist on acquisition)"],
        ["PUT",  "/radiology/studies/:id/status","Update acquisition status"],
        ["GET",  "/radiology/studies/:id/priors","Prior studies for same patient and anatomical region"],
    ],
    col_widths=[0.7, 2.5, 4.0])
add_note(doc, "pacsUrl should be a valid PACS viewer deeplink. Can be an OHIF viewer URL or DICOMweb link. Frontend will open in new tab.")

add_heading(doc, "9.3 Radiology Reports", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/radiology/reports",          "List reports. Filter: patientId, status, date"],
        ["GET",  "/radiology/reports/:id",      "Single report"],
        ["POST", "/radiology/reports",          "Create draft report. Body: { studyId, indication, technique }"],
        ["PUT",  "/radiology/reports/:id",      "Save/update report fields"],
        ["POST", "/radiology/reports/:id/sign", "Finalize & sign. Body: { radiologist }"],
        ["POST", "/radiology/reports/:id/addendum", "Append addendum. Body: { addendum, addendumBy }"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "9.4 Critical Findings", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/radiology/critical",                  "Pending critical findings. Filter: status"],
        ["POST", "/radiology/critical",                  "Flag critical finding. Body: CriticalFinding (without id)"],
        ["PUT",  "/radiology/critical/:id/notify",       "Record notification. Body: { notifiedTo, callbackNumber? }"],
        ["PUT",  "/radiology/critical/:id/acknowledge",  "Acknowledge. Body: { acknowledgedBy }"],
    ],
    col_widths=[0.7, 2.5, 4.0])
add_note(doc, "Critical findings must also trigger a WebSocket event: radiology.critical_finding and a CDSS recommendation (type: urgent_finding, severity: critical).")

add_heading(doc, "9.5 Modality Schedule (Slots)", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/radiology/schedule",           "Slots for a date range. Filter: modality, room, date"],
        ["POST", "/radiology/schedule",           "Create slot"],
        ["PUT",  "/radiology/schedule/:id",       "Update slot (book/cancel/block)"],
    ],
    col_widths=[0.7, 2.3, 4.2])

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 10 — PHARMACY
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "10. Module 6 — Pharmacy Portal")
add_para(doc, "Accessible by: pharmacist (primary), doctor (prescription view). Covers prescription queue, "
    "clinical verification, drug safety checks, dispensing, formulary, refills, substitutions, "
    "and pharmacist interventions.")

add_heading(doc, "10.1 Prescription Queue (Rx)", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/pharmacy/prescriptions",              "Rx queue. Filter: status, setting (inpatient/outpatient/discharge), priority, date"],
        ["GET",  "/pharmacy/prescriptions/:id",          "Single Rx with drug warnings"],
        ["PUT",  "/pharmacy/prescriptions/:id/verify",   "Clinical verification. Body: { verifiedBy, notes? }"],
        ["PUT",  "/pharmacy/prescriptions/:id/dispense", "Record dispensing. Body: DispenseRecord (without id)"],
        ["PUT",  "/pharmacy/prescriptions/:id/hold",     "Place on hold. Body: { reason }"],
        ["PUT",  "/pharmacy/prescriptions/:id/cancel",   "Cancel Rx"],
    ],
    col_widths=[0.7, 2.7, 3.8])

add_heading(doc, "PharmacyPrescription — Key Fields", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",              "string",    ""],
        ["patientId",       "string",    ""],
        ["patientName",     "string",    ""],
        ["mrn",             "string",    ""],
        ["medication",      "string",    "Brand name"],
        ["genericName",     "string",    ""],
        ["dosage",          "string",    "e.g. 500mg"],
        ["route",           "string",    "oral|iv|im|topical|inhaled|sublingual"],
        ["frequency",       "string",    "e.g. TID, BID, q6h"],
        ["quantity",        "number",    ""],
        ["refillsAllowed",  "number",    ""],
        ["refillsRemaining","number",    ""],
        ["setting",         "RxSetting", "inpatient|outpatient|discharge"],
        ["priority",        "Priority",  ""],
        ["status",          "RxStatus",  "ordered→pending-verification→verified→dispensing→dispensed|on-hold|cancelled|returned"],
        ["prescribedBy",    "string",    "ID/name of prescribing doctor"],
        ["warnings",        "DrugWarning[]","Auto-populated by CDSS/drug database"],
        ["allergies",       "string[]",  "Patient allergies from record — included for pharmacist display"],
    ],
    col_widths=[1.5, 1.2, 4.5])

add_heading(doc, "Drug Warning Object", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",            "string",          ""],
        ["type",          "WarningType",     "interaction|allergy|duplication|dose-range|renal|pregnancy|pediatric"],
        ["severity",      "WarningSeverity", "info|moderate|severe|contraindicated"],
        ["title",         "string",          "Short headline"],
        ["description",   "string",          "Clinical detail"],
        ["interactingDrug","string?",         "The other drug in an interaction"],
        ["overridable",   "boolean",         "Whether pharmacist can override"],
        ["overriddenBy",  "string?",         ""],
        ["overriddenAt",  "string?",         ""],
    ],
    col_widths=[1.4, 1.4, 4.4])

add_heading(doc, "10.2 Formulary", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/pharmacy/formulary",          "Search formulary. Query: q=, class=, status="],
        ["GET",  "/pharmacy/formulary/:id",      "Single formulary item"],
        ["PUT",  "/pharmacy/formulary/:id",      "Update item (admin/pharmacist)"],
        ["GET",  "/pharmacy/formulary/:id/stock","Current stock level for item"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "10.3 Medication Profiles", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET", "/patients/:id/medication-profile", "All active/recent medications for patient (used for CDSS + MAR)"],
    ],
    col_widths=[0.7, 2.8, 3.7])

add_heading(doc, "10.4 Interventions", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/pharmacy/interventions",         "List pharmacist interventions. Filter: outcome, date"],
        ["POST", "/pharmacy/interventions",         "Record intervention. Body: InterventionRecord (without id)"],
        ["PUT",  "/pharmacy/interventions/:id",     "Update outcome. Body: { outcome, prescriberResponse?, resolvedAt }"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "10.5 Refills", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/pharmacy/refills",      "Refill history. Filter: patientId, date"],
        ["POST", "/pharmacy/refills",      "Dispense refill. Body: RefillRecord (without id)"],
    ],
    col_widths=[0.7, 2.0, 4.5])

add_heading(doc, "10.6 Substitution Requests", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/pharmacy/substitutions",         "List substitution requests. Filter: status"],
        ["POST", "/pharmacy/substitutions",         "Request substitution"],
        ["PUT",  "/pharmacy/substitutions/:id",     "Approve/reject. Body: { status, approvedBy?, costSavings? }"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "10.7 Drug Safety Check (CDSS Integration)", level=2)
add_para(doc,
    "Whenever a prescription is created or a new medication is added, the frontend will call a drug safety "
    "check endpoint. The backend should integrate with a drug interaction database (e.g. Micromedex, Drugs.com API) "
    "and return warnings that will be embedded in the prescription object.")
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["POST", "/pharmacy/drug-safety-check", "Body: { patientId, medications: string[], newMedication: string }. Returns: DrugWarning[]"],
    ],
    col_widths=[0.7, 2.5, 4.0])

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 11 — BILLING
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "11. Module 7 — Billing / Revenue Cycle Portal")
add_para(doc, "Accessible by: billing_staff (primary). Covers patient accounts, invoices, "
    "insurance claims, payments, denials, and financial event timeline.")

add_heading(doc, "11.1 Patient Accounts", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET", "/billing/accounts",        "Patient account summaries. Filter: balance>0, overdue, search"],
        ["GET", "/billing/accounts/:patientId","Single patient account with balance breakdown"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "11.2 Invoices", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/billing/invoices",          "List invoices. Filter: status, patientId, date"],
        ["GET",  "/billing/invoices/:id",      "Single invoice with all charge items"],
        ["POST", "/billing/invoices",          "Create invoice from encounter/admission"],
        ["PUT",  "/billing/invoices/:id",      "Update invoice"],
        ["POST", "/billing/invoices/:id/send", "Mark as sent to patient/insurer"],
        ["POST", "/billing/invoices/:id/void", "Void invoice"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "Invoice Object — Key Fields", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",                   "string",              ""],
        ["patientId",            "string",              ""],
        ["encounterType",        "string",              "inpatient|outpatient|emergency|observation"],
        ["primaryDiagnosis",     "string",              "Free text + ICD-10 code"],
        ["chargeItems",          "ChargeItem[]",        "See ChargeItem below"],
        ["totalCharges",         "number",              "Sum of chargeItems"],
        ["insuranceBilled",      "number",              ""],
        ["insurancePaid",        "number",              ""],
        ["adjustments",          "number",              "Contractual & write-offs"],
        ["patientBalance",       "number",              "totalCharges - insurancePaid - adjustments - patientPaid"],
        ["status",               "BillingInvoiceStatus","draft|sent|billed_insurance|partial|unpaid|overdue|cleared|void"],
        ["insurancePlan",        "InsurancePlan?",      "Embedded payer info"],
    ],
    col_widths=[1.7, 1.5, 4.0])

add_heading(doc, "ChargeItem Object", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",                   "string",  ""],
        ["cptCode",              "string",  "CPT procedure code"],
        ["description",          "string",  ""],
        ["quantity",             "number",  ""],
        ["unitPrice",            "number",  ""],
        ["totalCharge",          "number",  "quantity × unitPrice"],
        ["diagnosisCodes",       "string[]","ICD-10 codes linking procedure to diagnosis"],
        ["serviceDate",          "string",  "YYYY-MM-DD"],
        ["department",           "string",  ""],
        ["provider",             "string",  "Rendering provider name/NPI"],
        ["modifier",             "string?", "CPT modifier (e.g. -25, -59)"],
        ["allowedAmount",        "number?", "Payer allowed amount (from EOB)"],
        ["adjustmentAmount",     "number?", ""],
        ["patientResponsibility","number?", "Copay + deductible portion"],
    ],
    col_widths=[1.7, 1.2, 4.3])

add_heading(doc, "11.3 Insurance Claims", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/billing/claims",               "List claims. Filter: status, payerId, date"],
        ["GET",  "/billing/claims/:id",           "Single claim"],
        ["POST", "/billing/claims",               "Submit claim from invoice"],
        ["PUT",  "/billing/claims/:id/status",    "Update claim status + EOB data. Body: { status, allowedAmount?, paidAmount?, eobDate? }"],
        ["POST", "/billing/claims/:id/resubmit",  "Resubmit after correction"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "11.4 Payments", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/billing/payments",         "Payment history. Filter: invoiceId, date, payer"],
        ["POST", "/billing/payments",         "Post payment (insurance or patient). Body: Payment (without id)"],
        ["PUT",  "/billing/payments/:id/void","Void payment. Body: { reason }"],
    ],
    col_widths=[0.7, 2.2, 4.3])

add_heading(doc, "11.5 Denial Management", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/billing/denials",              "Denial queue. Filter: status, reasonCode, date"],
        ["GET",  "/billing/denials/:id",          "Single denial"],
        ["POST", "/billing/denials/:id/appeal",   "Submit appeal. Body: { appealNotes, submittedAt }"],
        ["PUT",  "/billing/denials/:id",          "Update denial status/resolution"],
    ],
    col_widths=[0.7, 2.3, 4.2])

add_heading(doc, "11.6 Financial Timeline", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET", "/billing/accounts/:patientId/timeline", "All FinancialEvent records for patient, sorted by timestamp desc"],
    ],
    col_widths=[0.7, 2.8, 3.7])

add_heading(doc, "11.7 Billing Stats", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET", "/billing/stats", "Returns BillingStats: totalBilledToday, collectedToday, pendingInsurance, patientBalanceDue, pendingClaims, deniedClaims, overdue30Days, collectionRate"],
    ],
    col_widths=[0.7, 1.8, 4.7])

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 12 — ADMIN PORTAL
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "12. Module 8 — Admin Portal")
add_para(doc, "Accessible by: admin only. Covers user management, role permissions, departments, "
    "wards, beds, service catalogs, audit trail, and system settings.")

add_heading(doc, "12.1 User Management", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/admin/users",              "List users. Filter: role, status, department, search"],
        ["GET",  "/admin/users/:id",          "Single user"],
        ["POST", "/admin/users",              "Create user account. Body: AdminUser (without id)"],
        ["PUT",  "/admin/users/:id",          "Update user"],
        ["PUT",  "/admin/users/:id/status",   "Activate/suspend/deactivate. Body: { status }"],
        ["DELETE","/admin/users/:id",         "Permanently delete user (requires confirmation)"],
        ["POST", "/admin/users/:id/reset-password","Force password reset. Returns temp link"],
        ["GET",  "/admin/users/:id/activity", "Recent audit log entries for user"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "AdminUser Object", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",             "string",          ""],
        ["firstName",      "string",          ""],
        ["lastName",       "string",          ""],
        ["email",          "string",          "Unique"],
        ["phone",          "string?",         ""],
        ["role",           "AdminUserRole",   "One of the 9 role keys"],
        ["departmentId",   "string?",         "FK → AdminDepartment"],
        ["departmentName", "string?",         "Denormalised"],
        ["status",         "AdminUserStatus", "active|inactive|suspended|pending"],
        ["lastLogin",      "string?",         "ISO 8601"],
        ["createdAt",      "string",          "ISO 8601"],
        ["employeeId",     "string?",         "HR identifier"],
        ["specialization", "string?",         "e.g. Cardiology"],
        ["licenseNumber",  "string?",         "Medical license"],
    ],
    col_widths=[1.4, 1.4, 4.4])

add_heading(doc, "12.2 Role Permissions", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/admin/permissions",            "All role permission configurations"],
        ["PUT",  "/admin/permissions",            "Bulk update permissions. Body: RolePermission[]"],
    ],
    col_widths=[0.7, 2.2, 4.3])
add_para(doc, "PermissionAction: view | create | edit | delete | approve | export")
add_para(doc, "PermissionResource: patients | encounters | orders | prescriptions | lab_results | radiology_reports | invoices | claims | payments | users | departments | beds | catalogs | audit_logs | settings")

add_heading(doc, "12.3 Departments", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/admin/departments",         "List departments. Filter: type, status"],
        ["GET",  "/admin/departments/:id",     "Single department with staff + bed counts"],
        ["POST", "/admin/departments",         "Create department"],
        ["PUT",  "/admin/departments/:id",     "Update department"],
        ["PUT",  "/admin/departments/:id/status","Activate/deactivate"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "12.4 Wards & Beds", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/admin/wards",           "List wards. Filter: departmentId, type, status"],
        ["POST", "/admin/wards",           "Create ward"],
        ["PUT",  "/admin/wards/:id",       "Update ward"],
        ["GET",  "/admin/beds",            "All beds. Filter: wardId, status, type"],
        ["POST", "/admin/beds",            "Add bed"],
        ["PUT",  "/admin/beds/:id",        "Update bed (type, features, status)"],
        ["DELETE","/admin/beds/:id",       "Remove bed"],
    ],
    col_widths=[0.7, 2.2, 4.3])

add_heading(doc, "12.5 Service Catalogs", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/admin/catalogs/lab",              "Lab test catalog"],
        ["POST", "/admin/catalogs/lab",              "Add lab test"],
        ["PUT",  "/admin/catalogs/lab/:id",          "Update lab test"],
        ["GET",  "/admin/catalogs/radiology",        "Radiology exam catalog"],
        ["POST", "/admin/catalogs/radiology",        "Add exam"],
        ["PUT",  "/admin/catalogs/radiology/:id",    "Update exam"],
        ["GET",  "/admin/catalogs/services",         "General service catalog"],
        ["POST", "/admin/catalogs/services",         "Add service"],
        ["PUT",  "/admin/catalogs/services/:id",     "Update service"],
    ],
    col_widths=[0.7, 2.5, 4.0])

add_heading(doc, "12.6 Audit Trail", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET", "/admin/audit",   "Audit log. Filter: userId, role, action, resource, outcome, severity, date range. Paginated."],
    ],
    col_widths=[0.7, 1.8, 4.7])
add_heading(doc, "AuditLogEntry Object", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",         "string",       ""],
        ["timestamp",  "string",       "ISO 8601"],
        ["userId",     "string",       "FK → AdminUser"],
        ["userName",   "string",       ""],
        ["userRole",   "string",       "Role at time of action"],
        ["action",     "AuditAction",  "login|logout|create|read|update|delete|approve|reject|export|print|bulk_action|permission_change|password_reset|status_change"],
        ["resource",   "string",       "Resource type affected"],
        ["resourceId", "string?",      "ID of affected record"],
        ["details",    "string",       "Human-readable description"],
        ["ipAddress",  "string",       "Caller IP"],
        ["severity",   "string",       "info|warning|critical"],
        ["outcome",    "string",       "success|failure"],
        ["sessionId",  "string?",      ""],
    ],
    col_widths=[1.2, 1.3, 4.7])
add_note(doc, "Every mutating API action across ALL modules must append an AuditLogEntry. The audit log is immutable — no DELETE or UPDATE endpoint exists for it.")

add_heading(doc, "12.7 System Settings", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET", "/admin/settings",     "All system settings grouped by category"],
        ["PUT", "/admin/settings",     "Update multiple settings. Body: { key: value, ... }"],
    ],
    col_widths=[0.7, 2.0, 4.5])

add_heading(doc, "12.8 Admin Stats Dashboard", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET", "/admin/stats", "Returns AdminStats: totalUsers, activeUsers, totalDepartments, totalBeds, occupiedBeds, totalLabTests, totalRadiologyStudies, auditLogsToday, systemUptime"],
    ],
    col_widths=[0.7, 1.8, 4.7])

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 13 — CDSS
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "13. Module 9 — CDSS (Clinical Decision Support System)")
add_para(doc,
    "The CDSS is a hospital-wide, cross-module clinical intelligence layer. "
    "It is NOT a doctor-only feature — it serves nurses, pharmacists, lab technicians, radiologists, "
    "and doctors. The frontend subscribed to real-time CDSS alerts via WebSocket and can also "
    "poll REST endpoints.")

add_heading(doc, "13.1 Recommendation Endpoints", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/cdss/recommendations",                   "List all recommendations. Filter: patientId, severity, type, status, sourceModule, targetRole, dateFrom, dateTo"],
        ["GET",  "/cdss/recommendations/:id",               "Single recommendation (full object with explanation + evidenceSources)"],
        ["GET",  "/cdss/patients/:patientId/recommendations","All CDSS recommendations for a specific patient"],
        ["POST", "/cdss/recommendations",                   "Create new recommendation (from backend rule engine)"],
        ["PUT",  "/cdss/recommendations/:id/expire",        "Mark a recommendation expired"],
    ],
    col_widths=[0.7, 3.0, 3.5])
add_note(doc, "The /cdss/recommendations?targetRole=nurse endpoint should only return recommendations where the authenticated user's role is in the targetRoles array.")

add_heading(doc, "CDSSRecommendation Object", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",                    "string",              ""],
        ["patientId",             "string",              ""],
        ["patientName",           "string",              ""],
        ["patientMRN",            "string",              ""],
        ["encounterId",           "string?",             ""],
        ["sourceModule",          "CDSSSourceModule?",   "doctor|nursing|lab|pharmacy|radiology|emergency|surgery|system"],
        ["targetRoles",           "UserRole[]?",         "Which roles should see this alert"],
        ["type",                  "CDSSRecommendationType","One of 26 types (see below)"],
        ["severity",              "string",              "critical|warning|info"],
        ["status",                "string",              "active|acknowledged|overridden|dismissed|expired|followed"],
        ["title",                 "string",              "Short headline (≤80 chars)"],
        ["summary",               "string",              "1–3 sentence clinical summary"],
        ["explanation",           "CDSSExplanation",     "Detailed reasoning object (see below)"],
        ["evidenceSources",       "CDSSEvidenceSource[]","List of guideline/literature references"],
        ["generatedAt",           "string",              "ISO 8601"],
        ["expiresAt",             "string?",             "ISO 8601 — auto-expire if not acted on"],
        ["triggeredBy",           "string",              "What data event triggered this alert"],
        ["affectedMedications",   "string[]?",           "For drug safety alerts"],
        ["suggestedActions",      "string[]",            "Ordered list of recommended actions"],
        ["overrideReason",        "string?",             "Set when overridden"],
        ["overrideReasonCategory","string?",             ""],
        ["overriddenBy",          "string?",             ""],
        ["overriddenAt",          "string?",             "ISO 8601"],
        ["acknowledgedBy",        "string?",             ""],
        ["acknowledgedAt",        "string?",             "ISO 8601"],
        ["feedbackRating",        "1-5?",                "Clinician usefulness rating"],
        ["feedbackComment",       "string?",             ""],
    ],
    col_widths=[1.8, 1.5, 3.9])

add_heading(doc, "CDSSRecommendationType — Full Enum", level=3)
add_table(doc,
    ["Value", "Domain", "Description"],
    [
        ["drug_interaction",      "Medication",   "Known interaction between two active drugs"],
        ["allergy",               "Medication",   "Drug matches patient allergy"],
        ["dosage_warning",        "Medication",   "Dose outside safe range for patient weight/renal function"],
        ["duplicate_therapy",     "Medication",   "Two concurrent drugs from same class"],
        ["contraindication",      "Medication",   "Drug contraindicated given patient condition"],
        ["guideline",             "Order",        "Clinical practice guideline recommendation"],
        ["order_set",             "Order",        "Suggested order set for current diagnosis"],
        ["appropriateness_check", "Order",        "Imaging/procedure appropriateness alert"],
        ["diagnostic",            "Diagnostic",   "Suggested diagnostic workup"],
        ["abnormal_result",       "Diagnostic",   "Lab/vital outside normal range"],
        ["panic_value",           "Diagnostic",   "Critical lab value requiring immediate action"],
        ["delta_check",           "Diagnostic",   "Significant change from prior value"],
        ["critical_result",       "Diagnostic",   "Cross-module critical lab or imaging finding"],
        ["preventive",            "Preventive",   "Vaccination, screening, or preventive care due"],
        ["care_gap",              "Preventive",   "Evidence-based care gap identified"],
        ["follow_up_reminder",    "Preventive",   "Follow-up study or appointment recommended"],
        ["deterioration_alert",   "Nursing",      "Early warning score (NEWS2, MEWS) threshold reached"],
        ["overdue_task",          "Nursing",      "Scheduled nursing task past due"],
        ["risk_score",            "Nursing",      "Risk stratification (fall, pressure injury, VTE)"],
        ["urgent_finding",        "Radiology",    "Incidental or critical imaging finding"],
        ["triage_support",        "Emergency",    "ED triage decision support"],
        ["sepsis_alert",          "Emergency",    "Sepsis screening criteria met"],
        ["trauma_alert",          "Emergency",    "Trauma activation criteria triggered"],
        ["perioperative_warning", "Surgery",      "Pre/intra/post-operative safety warning"],
        ["checklist_gap",         "Surgery",      "Surgical safety checklist deviation"],
        ["care_plan_deviation",   "Cross",        "Patient care deviates from agreed plan"],
    ],
    col_widths=[1.7, 1.1, 4.4])

add_heading(doc, "CDSSExplanation Object", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["summary",         "string",          "1-paragraph plain language explanation"],
        ["reasoning",       "string[]",         "Bullet-point clinical reasoning steps"],
        ["clinicalInputs",  "CDSSClinicalInput[]","Key data points used: { label, value, flag? }"],
        ["limitations",     "string[]",         "Known limitations of this rule/model"],
        ["confidence",      "string",          "high|moderate|low"],
        ["confidenceScore", "number",          "0–100"],
        ["modelVersion",    "string?",         "Rule engine or AI model version identifier"],
    ],
    col_widths=[1.5, 1.5, 4.2])

add_heading(doc, "CDSSEvidenceSource Object", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",           "string",               ""],
        ["title",        "string",               "Full guideline/study title"],
        ["shortName",    "string",               "e.g. ADA 2026, ACC/AHA Guidelines"],
        ["sourceType",   "CDSSEvidenceSourceType","guideline|drug_database|literature|ai_model|ehr_pattern"],
        ["url",          "string?",              "Link to full guideline"],
        ["publishedYear","number?",              ""],
        ["evidenceGrade","string?",              "Level A / Grade 1B / Class I etc."],
        ["excerpt",      "string?",              "Key sentence from the source"],
    ],
    col_widths=[1.3, 1.5, 4.4])

add_heading(doc, "13.2 Override / Response Endpoints", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["POST", "/cdss/recommendations/:id/respond", "Clinician response. Body: CDSSOverridePayload (see below). Returns updated recommendation"],
        ["GET",  "/cdss/overrides",                   "Override audit trail. Filter: patientId, clinicianId, action, sourceModule, dateFrom, dateTo"],
        ["POST", "/cdss/recommendations/:id/feedback","Submit usefulness feedback. Body: { rating: 1-5, comment? }"],
    ],
    col_widths=[0.7, 2.7, 3.8])

add_heading(doc, "CDSSOverridePayload", level=3)
add_code(doc, "POST /cdss/recommendations/:id/respond")
add_code(doc, "Body: {")
add_code(doc, "  action:         CDSSOverrideAction,          // override|acknowledge|dismiss|follow")
add_code(doc, "  reasonCategory: CDSSOverrideReasonCategory,  // clinical_judgment|patient_preference|...")
add_code(doc, "  reason:         string,                      // free-text justification")
add_code(doc, "  notes?:         string,                      // additional notes")
add_code(doc, "  clinicianName:  string,")
add_code(doc, "  clinicianRole:  string,")
add_code(doc, "  sourceModule?:  CDSSSourceModule             // which UI context submitted this")
add_code(doc, "}")

add_heading(doc, "CDSSOverrideRecord — what gets persisted", level=3)
add_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id",                  "string",              ""],
        ["recommendationId",    "string",              "FK → CDSSRecommendation"],
        ["recommendationTitle", "string",              "Denormalised for audit display"],
        ["patientId",           "string",              ""],
        ["patientName",         "string",              ""],
        ["clinicianId",         "string",              "FK → User"],
        ["clinicianName",       "string",              ""],
        ["clinicianRole",       "string",              ""],
        ["action",              "CDSSOverrideAction",  "Which action was taken"],
        ["reasonCategory",      "string",              ""],
        ["reason",              "string",              ""],
        ["timestamp",           "string",              "ISO 8601 — when the action was recorded"],
        ["notes",               "string?",             ""],
        ["sourceModule",        "CDSSSourceModule?",   "Which portal submitted the override"],
    ],
    col_widths=[1.7, 1.5, 4.0])

add_heading(doc, "13.3 Stats", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET", "/cdss/stats", "Returns CDSSStats: totalActive, critical, warnings, info, overriddenToday, followedToday, acknowledgedToday, accuracyRate"],
    ],
    col_widths=[0.7, 1.8, 4.7])

add_heading(doc, "13.4 Real-Time CDSS Events (WebSocket)", level=2)
add_para(doc, "The frontend subscribes to a WebSocket channel on login. New CDSS recommendations are pushed immediately. "
    "The frontend toast-notifies the user for severity=critical.")
add_table(doc,
    ["Event Name", "Payload", "Description"],
    [
        ["cdss.new_recommendation", "CDSSRecommendation", "New alert generated"],
        ["cdss.recommendation_expired", "{ id: string }", "Alert expired"],
        ["cdss.recommendation_updated", "{ id, status }", "Status changed by another clinician"],
    ],
    col_widths=[2.2, 2.0, 3.0])

add_heading(doc, "13.5 Source Module Filtering Logic", level=2)
add_para(doc,
    "When the backend creates a CDSSRecommendation it MUST set sourceModule and targetRoles. "
    "The API must enforce that GET /cdss/recommendations only returns records where the authenticated "
    "user's role is in targetRoles. Legacy recommendations with no targetRoles set should be visible to all clinical roles.")
add_note(doc, "The frontend also has client-side adapter filtering as a UX layer, but the backend is the authoritative access control.")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 14 — PATIENT PORTAL
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "14. Module 10 — Patient Self-Service Portal")
add_para(doc, "Accessible by: patient role only. Read-heavy, limited write access.")

add_heading(doc, "14.1 Endpoints", level=2)
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/patients/me",                      "Authenticated patient's own profile"],
        ["GET",  "/patients/me/appointments",         "Upcoming + past appointments"],
        ["POST", "/patients/me/appointments",         "Request new appointment"],
        ["PUT",  "/patients/me/appointments/:id/cancel","Cancel appointment"],
        ["GET",  "/patients/me/prescriptions",        "Active prescriptions"],
        ["GET",  "/patients/me/lab-results",          "Released lab results only (status=final/verified)"],
        ["GET",  "/patients/me/vitals",               "Recent vitals (read-only)"],
        ["GET",  "/patients/me/encounters",           "Visit summaries (status=signed only)"],
        ["GET",  "/patients/me/invoices",             "Own billing statements"],
        ["POST", "/patients/me/invoices/:id/pay",     "Record patient payment attempt (redirects to payment gateway)"],
    ],
    col_widths=[0.7, 2.8, 3.7])
add_note(doc, "Patient portal must never expose draft encounters, unsigned notes, pending lab results, or other clinician-facing data.")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 15 — REAL-TIME & NOTIFICATIONS
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "15. Real-Time & Notification Requirements")

add_heading(doc, "15.1 WebSocket Events", level=2)
add_para(doc, "The frontend establishes a single authenticated WebSocket connection on login "
    "(ws://api.medhub.hospital/ws?token=JWT). Events are JSON objects with { event, payload }.")
add_table(doc,
    ["Event", "Subscribers", "Payload"],
    [
        ["cdss.new_recommendation",       "All clinical roles",         "CDSSRecommendation"],
        ["cdss.recommendation_updated",   "All clinical roles",         "{ id, status, severity }"],
        ["lab.critical_result",           "doctor, nurse, lab_tech",   "{ patientId, patientName, testName, value, flag }"],
        ["lab.result_released",           "doctor",                     "{ patientId, panelId, panelName }"],
        ["radiology.critical_finding",    "doctor, nurse, radiologist","CriticalFinding"],
        ["radiology.report_signed",       "doctor",                     "{ patientId, studyId, reportId, modality }"],
        ["adt.admission",                 "doctor, nurse, front_desk", "{ patientId, ward, bed }"],
        ["adt.discharge",                 "doctor, nurse, front_desk", "{ patientId }"],
        ["adt.bed_available",             "front_desk, nurse",          "{ bedId, ward }"],
        ["queue.ticket_called",           "patient, front_desk",        "{ ticketNo, window, patientId }"],
        ["pharmacy.rx_verified",          "doctor, nurse",              "{ patientId, medication }"],
        ["pharmacy.rx_dispensed",         "nurse",                      "{ patientId, medication, dispensedAt }"],
    ],
    col_widths=[2.2, 1.7, 3.3])

add_heading(doc, "15.2 Push Notifications (Optional)", level=2)
add_para(doc, "For mobile or installable PWA: implement Web Push (VAPID) or FCM for STAT/critical events.")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 16 — PAGINATION, FILTERING & SEARCH
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "16. Pagination, Filtering & Search Conventions")

add_heading(doc, "16.1 List Endpoint Query Parameters", level=2)
add_table(doc,
    ["Parameter", "Type", "Description"],
    [
        ["page",     "number",  "1-based page number (default: 1)"],
        ["limit",    "number",  "Records per page (default: 20, max: 100)"],
        ["sortBy",   "string",  "Field name to sort by"],
        ["sortOrder","string",  "asc | desc (default: desc for dates, asc for names)"],
        ["search",   "string",  "Free-text search across relevant indexed fields"],
        ["dateFrom", "string",  "ISO 8601 date — filter start"],
        ["dateTo",   "string",  "ISO 8601 date — filter end"],
        ["status",   "string",  "Status filter (specific to resource)"],
    ],
    col_widths=[1.2, 0.8, 5.2])

add_heading(doc, "16.2 Paginated Response Envelope", level=2)
add_code(doc, "{")
add_code(doc, "  data:       T[],        // array of items for current page")
add_code(doc, "  total:      number,     // total matching records")
add_code(doc, "  page:       number,     // current page")
add_code(doc, "  limit:      number,     // page size")
add_code(doc, "  totalPages: number      // Math.ceil(total / limit)")
add_code(doc, "}")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 17 — ERROR RESPONSE FORMAT
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "17. Error Response Format")
add_para(doc, "All error responses must use the following standardised body regardless of HTTP status code:")
add_code(doc, "{")
add_code(doc, "  error: {")
add_code(doc, "    code:    string,    // Machine-readable (e.g. PATIENT_NOT_FOUND, VALIDATION_ERROR)")
add_code(doc, "    message: string,    // Human-readable description")
add_code(doc, "    details: object?    // Optional extra context (e.g. validation field errors)")
add_code(doc, "  }")
add_code(doc, "}")
add_para(doc, "")
add_table(doc,
    ["HTTP Status", "When to Use"],
    [
        ["200 OK",              "Successful GET, PUT, PATCH"],
        ["201 Created",         "Successful POST creating a resource"],
        ["204 No Content",      "Successful DELETE"],
        ["400 Bad Request",     "Validation error, malformed JSON"],
        ["401 Unauthorized",    "Missing or invalid JWT"],
        ["403 Forbidden",       "Valid JWT but insufficient role/permission"],
        ["404 Not Found",       "Resource ID does not exist"],
        ["409 Conflict",        "Duplicate resource (e.g. duplicate MRN, duplicate claim)"],
        ["422 Unprocessable",   "Business logic error (e.g. discharging a non-admitted patient)"],
        ["500 Internal Error",  "Unexpected server error — should be logged"],
    ],
    col_widths=[1.5, 5.7])

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 18 — AUDIT LOGGING
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "18. Audit Logging Requirements")
add_para(doc,
    "Every API action that creates, reads (PHI), updates, or deletes data must be captured in the audit log. "
    "The audit log is write-only and must never be modified or deleted.")
add_table(doc,
    ["Action Type", "Examples", "Severity"],
    [
        ["login / logout",   "Successful or failed logins",                   "info"],
        ["create",           "New patient, new Rx, new admission",             "info"],
        ["read (PHI)",       "Viewing patient chart, lab results",             "info"],
        ["update",           "Editing encounter, changing user role",          "info"],
        ["delete",           "Removing a bed, cancelling an order",            "warning"],
        ["approve",          "Signing encounter, releasing lab result",        "info"],
        ["export / print",   "Exporting patient data, printing report",        "warning"],
        ["permission_change","Modifying RBAC roles",                           "critical"],
        ["password_reset",   "Admin-forced password reset",                    "warning"],
        ["CDSS override",    "Clinician overrides a safety alert",             "warning"],
    ],
    col_widths=[1.8, 3.1, 1.0])

add_note(doc, "HIPAA compliance: retain audit logs for a minimum of 6 years. Logs must include IP address and session ID.")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 19 — FILE UPLOADS
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "19. File / Document Upload")
add_table(doc,
    ["Use Case", "Endpoint", "Notes"],
    [
        ["Patient avatar",              "PUT /patients/:id/avatar",          "multipart/form-data, max 2MB, JPEG/PNG"],
        ["Consent document (PDF)",      "POST /consents/:id/upload",         "PDF only, max 10MB, store in S3/object storage"],
        ["Radiology DICOM",             "POST /radiology/studies/:id/images","DICOM files; recommend separate PACS upload API"],
        ["Lab attachment",              "POST /lab/reports/:id/attachment",  "PDF results from external analyzers"],
        ["Insurance card scan",         "POST /patients/:id/insurance/card", "JPEG/PNG/PDF max 5MB"],
    ],
    col_widths=[1.8, 2.4, 3.0])
add_para(doc, "All file upload responses must return: { fileUrl: string, fileId: string, uploadedAt: string }")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 20 — ENDPOINT CHECKLIST
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "20. Summary Endpoint Checklist")
add_para(doc, "Complete list of all endpoints the frontend consumes, grouped by module.")

modules_checklist = [
    ("Auth", [
        "POST /auth/login",
        "POST /auth/logout",
        "POST /auth/refresh",
        "GET  /auth/me",
        "PUT  /auth/me/password",
    ]),
    ("Patients / ADT", [
        "GET    /patients",
        "GET    /patients/:id",
        "POST   /patients",
        "PUT    /patients/:id",
        "GET    /patients/search",
        "GET    /patients/:id/duplicates",
        "POST   /patients/merge",
        "GET    /admissions",
        "POST   /admissions",
        "PUT    /admissions/:id/status",
        "POST   /admissions/:id/discharge",
        "POST   /admissions/:id/transfer",
        "GET    /beds",
        "PUT    /beds/:id/status",
        "GET    /wards",
        "GET    /queue",
        "POST   /queue",
        "PUT    /queue/:id/status",
        "GET    /patients/:id/consents",
        "POST   /patients/:id/consents",
        "PUT    /consents/:id/sign",
    ]),
    ("Doctor", [
        "GET    /appointments",
        "POST   /appointments",
        "PUT    /appointments/:id",
        "PUT    /appointments/:id/status",
        "DELETE /appointments/:id",
        "GET    /patients/:id/encounters",
        "GET    /encounters/:id",
        "POST   /encounters",
        "PUT    /encounters/:id",
        "POST   /encounters/:id/sign",
        "POST   /encounters/:id/amend",
        "GET    /patients/:id/orders",
        "POST   /orders",
        "PUT    /orders/:id/status",
        "DELETE /orders/:id",
        "GET    /patients/:id/prescriptions",
        "POST   /prescriptions",
        "PUT    /prescriptions/:id/status",
        "GET    /patients/:id/diagnoses",
        "POST   /diagnoses",
        "PUT    /diagnoses/:id",
        "GET    /icd10/search",
        "GET    /patients/:id/referrals",
        "POST   /referrals",
        "PUT    /referrals/:id/status",
        "GET    /doctors/:id/results",
        "PUT    /results/:id/review",
    ]),
    ("Nurse", [
        "GET  /wards/:wardId/patients",
        "GET  /patients/:id/vitals",
        "POST /patients/:id/vitals",
        "GET  /patients/:id/vitals/latest",
        "GET  /patients/:id/io",
        "POST /patients/:id/io",
        "GET  /patients/:id/pain",
        "POST /patients/:id/pain",
        "GET  /patients/:id/mar",
        "PUT  /mar/:id/administer",
        "PUT  /mar/:id/status",
        "GET  /patients/:id/nursing-notes",
        "POST /patients/:id/nursing-notes",
        "GET  /nurses/:id/tasks",
        "PUT  /tasks/:id/complete",
        "GET  /patients/:id/wounds",
        "POST /patients/:id/wounds",
        "GET  /handoffs",
        "POST /handoffs",
        "GET  /patients/:id/discharge-checklist",
        "PUT  /discharge-checklist/:id/complete",
    ]),
    ("Lab (LIS)", [
        "GET  /lab/worklist",
        "GET  /specimens",
        "PUT  /specimens/:id/status",
        "POST /specimens/:id/reject",
        "POST /specimens/recollect",
        "GET  /accessions",
        "POST /accessions",
        "GET  /analyzers/queue",
        "PUT  /analyzers/queue/:id/status",
        "GET  /lab/panels",
        "GET  /lab/panels/:id",
        "POST /lab/panels/:id/results",
        "PUT  /lab/panels/:id/verify",
        "GET  /lab/reports",
        "GET  /lab/reports/:id",
        "PUT  /lab/reports/:id/release",
        "GET  /lab/critical",
        "POST /lab/critical/:resultId/notify",
        "PUT  /lab/critical/:resultId/acknowledge",
    ]),
    ("Radiology (RIS)", [
        "GET  /radiology/orders",
        "POST /radiology/orders",
        "PUT  /radiology/orders/:id/protocol",
        "PUT  /radiology/orders/:id/schedule",
        "PUT  /radiology/orders/:id/assign",
        "PUT  /radiology/orders/:id/cancel",
        "GET  /radiology/studies",
        "GET  /radiology/studies/:id",
        "PUT  /radiology/studies/:id/status",
        "GET  /radiology/studies/:id/priors",
        "GET  /radiology/reports",
        "GET  /radiology/reports/:id",
        "POST /radiology/reports",
        "PUT  /radiology/reports/:id",
        "POST /radiology/reports/:id/sign",
        "POST /radiology/reports/:id/addendum",
        "GET  /radiology/critical",
        "POST /radiology/critical",
        "PUT  /radiology/critical/:id/notify",
        "PUT  /radiology/critical/:id/acknowledge",
        "GET  /radiology/schedule",
        "POST /radiology/schedule",
        "PUT  /radiology/schedule/:id",
    ]),
    ("Pharmacy", [
        "GET  /pharmacy/prescriptions",
        "GET  /pharmacy/prescriptions/:id",
        "PUT  /pharmacy/prescriptions/:id/verify",
        "PUT  /pharmacy/prescriptions/:id/dispense",
        "PUT  /pharmacy/prescriptions/:id/hold",
        "PUT  /pharmacy/prescriptions/:id/cancel",
        "GET  /pharmacy/formulary",
        "GET  /pharmacy/formulary/:id",
        "GET  /patients/:id/medication-profile",
        "GET  /pharmacy/interventions",
        "POST /pharmacy/interventions",
        "PUT  /pharmacy/interventions/:id",
        "GET  /pharmacy/refills",
        "POST /pharmacy/refills",
        "GET  /pharmacy/substitutions",
        "POST /pharmacy/substitutions",
        "PUT  /pharmacy/substitutions/:id",
        "POST /pharmacy/drug-safety-check",
    ]),
    ("Billing", [
        "GET  /billing/accounts",
        "GET  /billing/accounts/:patientId",
        "GET  /billing/accounts/:patientId/timeline",
        "GET  /billing/invoices",
        "GET  /billing/invoices/:id",
        "POST /billing/invoices",
        "PUT  /billing/invoices/:id",
        "POST /billing/invoices/:id/send",
        "POST /billing/invoices/:id/void",
        "GET  /billing/claims",
        "GET  /billing/claims/:id",
        "POST /billing/claims",
        "PUT  /billing/claims/:id/status",
        "POST /billing/claims/:id/resubmit",
        "GET  /billing/payments",
        "POST /billing/payments",
        "PUT  /billing/payments/:id/void",
        "GET  /billing/denials",
        "GET  /billing/denials/:id",
        "POST /billing/denials/:id/appeal",
        "PUT  /billing/denials/:id",
        "GET  /billing/stats",
    ]),
    ("Admin", [
        "GET    /admin/users",
        "GET    /admin/users/:id",
        "POST   /admin/users",
        "PUT    /admin/users/:id",
        "PUT    /admin/users/:id/status",
        "DELETE /admin/users/:id",
        "POST   /admin/users/:id/reset-password",
        "GET    /admin/permissions",
        "PUT    /admin/permissions",
        "GET    /admin/departments",
        "POST   /admin/departments",
        "PUT    /admin/departments/:id",
        "GET    /admin/wards",
        "POST   /admin/wards",
        "GET    /admin/beds",
        "POST   /admin/beds",
        "PUT    /admin/beds/:id",
        "DELETE /admin/beds/:id",
        "GET    /admin/catalogs/lab",
        "POST   /admin/catalogs/lab",
        "GET    /admin/catalogs/radiology",
        "POST   /admin/catalogs/radiology",
        "GET    /admin/catalogs/services",
        "POST   /admin/catalogs/services",
        "GET    /admin/audit",
        "GET    /admin/settings",
        "PUT    /admin/settings",
        "GET    /admin/stats",
    ]),
    ("CDSS", [
        "GET  /cdss/recommendations",
        "GET  /cdss/recommendations/:id",
        "GET  /cdss/patients/:patientId/recommendations",
        "POST /cdss/recommendations",
        "PUT  /cdss/recommendations/:id/expire",
        "POST /cdss/recommendations/:id/respond",
        "GET  /cdss/overrides",
        "POST /cdss/recommendations/:id/feedback",
        "GET  /cdss/stats",
        "WS   cdss.new_recommendation",
        "WS   cdss.recommendation_updated",
    ]),
    ("Patient Portal", [
        "GET  /patients/me",
        "GET  /patients/me/appointments",
        "POST /patients/me/appointments",
        "PUT  /patients/me/appointments/:id/cancel",
        "GET  /patients/me/prescriptions",
        "GET  /patients/me/lab-results",
        "GET  /patients/me/vitals",
        "GET  /patients/me/encounters",
        "GET  /patients/me/invoices",
        "POST /patients/me/invoices/:id/pay",
    ]),
]

for section_name, endpoints in modules_checklist:
    add_heading(doc, section_name, level=2)
    for ep in endpoints:
        add_bullet(doc, ep)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# CLOSING NOTE
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "Notes & Implementation Guidance")
add_para(doc,
    "This document was generated directly from the frontend TypeScript types, Zustand stores, "
    "Zod validation schemas, and component logic. Every field described here is actively consumed "
    "by the frontend.", size=10)
add_para(doc, "")
add_bullet(doc, "All money/financial values should be in the smallest currency unit (e.g. USD cents) OR a float with 2 decimal places — agree with frontend team before implementation.")
add_bullet(doc, "IDs should be UUID strings. The frontend uses string IDs throughout.")
add_bullet(doc, "Timestamps should always be UTC ISO 8601. The frontend uses date-fns for local display conversion.")
add_bullet(doc, "Any field marked with ? in the type definitions is optional — the frontend handles undefined gracefully.")
add_bullet(doc, "The backend should NOT send placeholder/mock data in production. Every field the frontend reads must be backed by real persistence.")
add_bullet(doc, "CORS: allow the frontend origin. In production, restrict to the deployed domain.")
add_bullet(doc, "Rate limiting: recommended 100 req/min per user for general endpoints; 10 req/min for auth.")
add_bullet(doc, "The frontend uses TanStack Query for caching — stale time is typically 30s–5min per resource type. The backend must set Cache-Control headers appropriately.")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"— End of Specification — Generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')} —")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run.italic = True

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────

out_path = "/Users/khaledsamy/Graduation Project/frontend copy/VirtualHospital_Backend_Spec.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
